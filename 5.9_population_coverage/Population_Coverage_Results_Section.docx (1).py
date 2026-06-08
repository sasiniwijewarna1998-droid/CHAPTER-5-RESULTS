from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ============================================================
# PLACEMENT INSTRUCTIONS
# ============================================================
p = doc.add_paragraph()
run = p.add_run("WHERE TO INSERT THIS SECTION IN YOUR THESIS")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(180, 0, 0)

doc.add_paragraph("")

# Instruction box for Results
p = doc.add_paragraph()
run = p.add_run("RESULTS (Chapter 5):")
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 180)
p = doc.add_paragraph()
run = p.add_run(
    "Insert as Section 5.4.4 immediately AFTER Section 5.4.3 (Helper T Lymphocyte Epitope Selection, "
    "paragraph 1437, Table 5.10) and BEFORE Section 5.5 (Multi-Epitope Vaccine Assembly, paragraph 1448).\n\n"
    "This placement is logical because population coverage analysis is the final validation step "
    "of your epitope selection pipeline, confirming that the selected CTL and HTL epitopes collectively "
    "cover the target populations before proceeding to vaccine construct assembly.\n\n"
    "Numbering: Use Section 5.4.4 as a subsection under 5.4 (Epitope Screening and Selection). "
    "This avoids renumbering all subsequent sections (5.5 through 5.10). The new tables and figures "
    "are numbered 5.10a through 5.10d and Figure 5.2a and 5.2b to slot between existing elements."
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0, 0, 180)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("DISCUSSION (Chapter 6):")
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)
p = doc.add_paragraph()
run = p.add_run(
    "Insert the discussion paragraph AFTER the paragraph on immunological safety screening "
    "(paragraph 1839, ending with '...favorable preliminary safety profiles...') and BEFORE "
    "the paragraph on multi-epitope vaccine construction (paragraph 1840, starting with "
    "'The multi-epitope vaccine construction phase...').\n\n"
    "This placement follows your Discussion's logical flow: epitope screening results are discussed, "
    "then population coverage validates the epitope selection, then vaccine construction is discussed."
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_page_break()

# ============================================================
# RESULTS SECTION
# ============================================================
h = doc.add_heading('5.4.4 Population Coverage Analysis', level=4)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Intro paragraph
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Population coverage analysis was performed using the IEDB Population Coverage Tool (v3.0.2) "
    "to evaluate the proportion of individuals within target populations expected to mount an immune "
    "response against the selected epitope ensemble. The analysis incorporated 28 HLA Class I alleles "
    "representing 12 MHC-I supertypes (A1, A2, A3, A24, A26, B7, B8, B27, B39, B44, B58, and B62) "
    "and 23 HLA Class II alleles spanning the DRB1, DRB3, DRB4, DRB5, DPA1, and DQA1 loci. Coverage "
    "was assessed for four target populations: World, South Asia, Southeast Asia, and Sri Lanka. India "
    "was additionally included as a South Asian comparator with more comprehensive HLA typing data in "
    "the IEDB Allele Frequency Database."
)

# ------ HLA ALLELE TABLE ------
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Table 5.10a ")
run.bold = True
p.add_run("Major HLA alleles included in the population coverage analysis. A total of 51 unique HLA "
          "alleles were derived from the restricting elements of the 15 selected CTL and 15 selected HTL epitopes.")
p.paragraph_format.space_after = Pt(6)

# HLA Table
hla_data = {
    'Class I (HLA-A)': [
        'HLA-A*01:01', 'HLA-A*02:01', 'HLA-A*02:03', 'HLA-A*02:06',
        'HLA-A*03:01', 'HLA-A*11:01', 'HLA-A*23:01', 'HLA-A*24:02',
        'HLA-A*25:01', 'HLA-A*26:01', 'HLA-A*31:01', 'HLA-A*68:01'
    ],
    'Class I (HLA-B)': [
        'HLA-B*07:02', 'HLA-B*08:01', 'HLA-B*14:02', 'HLA-B*15:01',
        'HLA-B*27:05', 'HLA-B*35:01', 'HLA-B*38:01', 'HLA-B*39:01',
        'HLA-B*40:01', 'HLA-B*44:02', 'HLA-B*44:03', 'HLA-B*46:01',
        'HLA-B*51:01', 'HLA-B*53:01', 'HLA-B*57:01', 'HLA-B*58:01'
    ],
    'Class II (HLA-DRB)': [
        'HLA-DRB1*01:01', 'HLA-DRB1*03:01', 'HLA-DRB1*04:01', 'HLA-DRB1*04:05',
        'HLA-DRB1*07:01', 'HLA-DRB1*08:02', 'HLA-DRB1*09:01', 'HLA-DRB1*11:01',
        'HLA-DRB1*12:01', 'HLA-DRB1*13:02', 'HLA-DRB1*15:01',
        'HLA-DRB3*01:01', 'HLA-DRB3*02:02', 'HLA-DRB4*01:01', 'HLA-DRB5*01:01'
    ],
    'Class II (HLA-DP/DQ)': [
        'HLA-DPA1*01:03', 'HLA-DPA1*02:01', 'HLA-DPA1*03:01',
        'HLA-DQA1*01:01', 'HLA-DQA1*01:02', 'HLA-DQA1*03:01',
        'HLA-DQA1*04:01', 'HLA-DQA1*05:01'
    ]
}

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
hdr = table.rows[0].cells
for idx, text in enumerate(['HLA Class', 'Locus Group', 'Alleles']):
    hdr[idx].text = text
    for p in hdr[idx].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

for cls_name, alleles in hla_data.items():
    row = table.add_row().cells
    if 'Class I' in cls_name:
        row[0].text = 'Class I'
    else:
        row[0].text = 'Class II'
    row[1].text = cls_name.split('(')[1].rstrip(')')
    row[2].text = ', '.join(alleles)
    for cell in row:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

doc.add_paragraph("")

# ------ 5.4.4.1 CLASS I RESULTS ------
h = doc.add_heading('5.4.4.1 MHC Class I (CTL Epitope) Coverage', level=5)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "The 15 selected CTL epitopes collectively achieved high population coverage across the global "
    "population and major dengue-endemic regions (Table 5.10b). The world population exhibited 98.20% "
    "coverage, with an average of 10.13 epitope-HLA interactions per individual and a PC90 value of "
    "4.24, indicating that 90% of the world population would recognize at least four of the selected "
    "CTL epitopes. Southeast Asia demonstrated 96.10% coverage (average hits = 9.39; PC90 = 2.96), "
    "while South Asia reached 91.18% coverage (average hits = 6.97; PC90 = 2.09). India, as a South "
    "Asian comparator with more extensive HLA typing studies, showed 87.51% Class I coverage (average "
    "hits = 6.28; PC90 = 1.60)."
)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Sri Lanka exhibited a comparatively lower Class I coverage of 52.39% (average hits = 2.39; PC90 "
    "= 0.42). However, this result is attributable to the limited availability of Sri Lankan HLA "
    "Class I allele frequency data within the IEDB Allele Frequency Database, rather than a genuine "
    "deficiency in epitope-HLA coverage. The IEDB database contains HLA frequency data from a restricted "
    "number of Sri Lankan typing studies, resulting in an incomplete representation of the Sri Lankan "
    "HLA repertoire. This interpretation is supported by the substantially higher coverage observed "
    "for the broader South Asian region (91.18%), of which Sri Lanka is a constituent population."
)

# Table 5.10b: Class I Coverage
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Table 5.10b ")
run.bold = True
p.add_run("MHC Class I population coverage of the 15 selected CTL epitopes across target populations.")
p.paragraph_format.space_after = Pt(6)

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
hdr = table.rows[0].cells
for idx, text in enumerate(['Population', 'Coverage (%)', 'Average Hits', 'PC90']):
    hdr[idx].text = text
    for p in hdr[idx].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

class_i_data = [
    ('World', '98.20', '10.13', '4.24'),
    ('South Asia', '91.18', '6.97', '2.09'),
    ('Southeast Asia', '96.10', '9.39', '2.96'),
    ('Sri Lanka', '52.39*', '2.39', '0.42'),
    ('India', '87.51', '6.28', '1.60'),
]
for pop, cov, hits, pc in class_i_data:
    row = table.add_row().cells
    row[0].text = pop
    row[1].text = cov
    row[2].text = hits
    row[3].text = pc
    for cell in row:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

p = doc.add_paragraph()
p.add_run("*Limited HLA Class I frequency data available in the IEDB Allele Frequency Database for Sri Lanka.").italic = True
p.runs[0].font.size = Pt(9)
p = doc.add_paragraph()
p.add_run("PC90 = minimum number of epitope-HLA interactions recognized by 90% of the population; Average Hits = mean number of epitope-HLA combinations recognized per individual.").italic = True
p.runs[0].font.size = Pt(9)

doc.add_paragraph("")

# ------ 5.4.4.2 CLASS II RESULTS ------
h = doc.add_heading('5.4.4.2 MHC Class II (HTL Epitope) Coverage', level=5)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "The 15 selected HTL epitopes achieved exceptionally high MHC Class II coverage across all "
    "populations for which HLA Class II frequency data was available (Table 5.10c). Global coverage "
    "reached 99.68%, with an average of 49.87 epitope-HLA interactions per individual and a PC90 "
    "of 34.79. South Asia demonstrated the highest regional coverage at 99.74% (average hits = "
    "46.11; PC90 = 32.47), comparable to India alone at 99.74% (average hits = 46.00; PC90 = 32.39). "
    "Southeast Asia showed 94.53% coverage (average hits = 27.81; PC90 = 17.25)."
)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "HLA Class II allele frequency data for Sri Lanka was not available in the IEDB Allele Frequency "
    "Database at the time of analysis; therefore, population-specific Class II coverage could not be "
    "independently computed for Sri Lanka. Given that the broader South Asian population, which "
    "encompasses Sri Lanka, achieved 99.74% Class II coverage, a comparably high coverage for the "
    "Sri Lankan population can be reasonably inferred."
)

# Table 5.10c: Class II Coverage
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Table 5.10c ")
run.bold = True
p.add_run("MHC Class II population coverage of the 15 selected HTL epitopes across target populations.")
p.paragraph_format.space_after = Pt(6)

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
hdr = table.rows[0].cells
for idx, text in enumerate(['Population', 'Coverage (%)', 'Average Hits', 'PC90']):
    hdr[idx].text = text
    for p in hdr[idx].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

class_ii_data = [
    ('World', '99.68', '49.87', '34.79'),
    ('South Asia', '99.74', '46.11', '32.47'),
    ('Southeast Asia', '94.53', '27.81', '17.25'),
    ('Sri Lanka', 'N/A**', 'N/A**', 'N/A**'),
    ('India', '99.74', '46.00', '32.39'),
]
for pop, cov, hits, pc in class_ii_data:
    row = table.add_row().cells
    row[0].text = pop
    row[1].text = cov
    row[2].text = hits
    row[3].text = pc
    for cell in row:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

p = doc.add_paragraph()
p.add_run("**No HLA Class II allele frequency data available for Sri Lanka in the IEDB Allele Frequency Database.").italic = True
p.runs[0].font.size = Pt(9)

doc.add_paragraph("")

# ------ 5.4.4.3 COMBINED RESULTS ------
h = doc.add_heading('5.4.4.3 Combined (Class I + Class II) Population Coverage', level=5)
for run in h.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "When MHC Class I and Class II coverage were combined, near-universal population coverage was "
    "achieved for all regions with comprehensive HLA frequency data (Table 5.10d, Figure 5.2a). The "
    "combined global coverage was 99.99% (average hits = 60.01; PC90 = 37.66), indicating that "
    "virtually the entire world population would recognize at least one epitope from the selected "
    "ensemble, with 90% of individuals recognizing 37 or more epitope-HLA combinations. South Asia "
    "reached 99.98% combined coverage (average hits = 53.07; PC90 = 32.63), and Southeast Asia "
    "achieved 99.79% (average hits = 37.20; PC90 = 18.86). India alone showed 99.97% combined "
    "coverage (average hits = 52.28; PC90 = 32.30)."
)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "The combined coverage for Sri Lanka (52.39%) reflects Class I coverage only, as Class II "
    "frequency data was unavailable. This value should therefore not be interpreted as the actual "
    "combined immunological coverage for the Sri Lankan population. Based on the regional South "
    "Asian coverage of 99.98%, the true combined coverage for Sri Lanka is expected to substantially "
    "exceed the reported Class I-only estimate."
)

# Table 5.10d: Combined Coverage
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Table 5.10d ")
run.bold = True
p.add_run("Combined (MHC Class I + Class II) population coverage of the selected CTL and HTL epitope "
          "ensemble across target populations.")
p.paragraph_format.space_after = Pt(6)

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
hdr = table.rows[0].cells
for idx, text in enumerate(['Population', 'Coverage (%)', 'Average Hits', 'PC90']):
    hdr[idx].text = text
    for p in hdr[idx].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

combined_data = [
    ('World', '99.99', '60.01', '37.66'),
    ('South Asia', '99.98', '53.07', '32.63'),
    ('Southeast Asia', '99.79', '37.20', '18.86'),
    ('Sri Lanka', '52.39*', '2.39', '0.42'),
    ('India', '99.97', '52.28', '32.30'),
]
for pop, cov, hits, pc in combined_data:
    row = table.add_row().cells
    row[0].text = pop
    row[1].text = cov
    row[2].text = hits
    row[3].text = pc
    for cell in row:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

p = doc.add_paragraph()
p.add_run("*Combined coverage reflects Class I only due to absence of Class II HLA frequency data for Sri Lanka in the IEDB database.").italic = True
p.runs[0].font.size = Pt(9)

doc.add_paragraph("")

# ------ FIGURES ------
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.add_run("The population coverage distribution across MHC Class I, Class II, and combined analyses "
          "is illustrated in Figure 5.2a. The cumulative coverage curves generated by the IEDB Population "
          "Coverage Tool for each population and analysis type are presented in Figure 5.2b.")
p.paragraph_format.line_spacing = 1.5

# Figure 5.2a
doc.add_paragraph("")
if os.path.exists('/workspace/session/Population_Coverage_BarChart.png'):
    doc.add_picture('/workspace/session/Population_Coverage_BarChart.png', width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Figure 5.2a ")
run.bold = True
p.add_run(
    "Population coverage of the selected CTL and HTL epitope ensemble across target populations. "
    "(A) MHC Class I coverage achieved by 15 CTL epitopes restricted by 28 HLA-I alleles. "
    "(B) MHC Class II coverage achieved by 15 HTL epitopes restricted by 23 HLA-II alleles; "
    "Sri Lanka Class II data was unavailable in the IEDB database. "
    "(C) Combined Class I and Class II coverage; Sri Lanka value reflects Class I only. "
    "The dashed red line indicates the 90% coverage threshold. Analysis was performed using the "
    "IEDB Population Coverage Tool (v3.0.2)."
)
for run in p.runs:
    run.font.size = Pt(10)

doc.add_page_break()

# Figure 5.2b
if os.path.exists('/workspace/session/Population_Coverage_IEDB_Plots.png'):
    doc.add_picture('/workspace/session/Population_Coverage_IEDB_Plots.png', width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Figure 5.2b ")
run.bold = True
p.add_run(
    "Cumulative population coverage distribution curves generated by the IEDB Population Coverage "
    "Tool for MHC Class I (top row), MHC Class II (middle row), and combined (bottom row) analyses "
    "across World, South Asia, Southeast Asia, and Sri Lanka populations. Each panel displays the "
    "percentage of individuals (left y-axis, blue bars) and cumulative percentage of population "
    "coverage (right y-axis, green line) as a function of the number of epitope-HLA combinations "
    "recognized. The red horizontal line indicates the 90% coverage threshold. No Class II HLA "
    "frequency data was available for Sri Lanka in the IEDB database."
)
for run in p.runs:
    run.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# DISCUSSION PARAGRAPH
# ============================================================
h = doc.add_heading('DISCUSSION PARAGRAPH', level=3)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 128, 0)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
run = p.add_run(
    "INSERT LOCATION: After the paragraph ending with '...favorable preliminary safety profiles...' "
    "(paragraph 1839) and BEFORE 'The multi-epitope vaccine construction phase...' (paragraph 1840) "
    "in Chapter 6."
)
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph("")

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "The population coverage analysis performed using the IEDB Population Coverage Tool confirmed "
    "that the selected epitope ensemble provided extensive HLA coverage across the primary target "
    "populations. The combined MHC Class I and Class II coverage reached 99.99% for the global "
    "population, 99.98% for South Asia, and 99.79% for Southeast Asia, the two principal regions of "
    "dengue endemicity relevant to this study. These coverage values substantially exceed the commonly "
    "accepted threshold of 90% population coverage for multi-epitope vaccine candidates (Bui et al., "
    "2006). The high PC90 values observed (37.66 for the world population and 32.63 for South Asia) "
    "further indicate that 90% of individuals in these populations would recognize more than 30 "
    "epitope-HLA combinations from the selected ensemble, reflecting substantial redundancy in "
    "immune recognition. The Class I coverage of 98.20% globally and 96.10% in Southeast Asia "
    "supports the capacity of the selected CTL epitopes to elicit broad CD8+ cytotoxic responses, "
    "while the Class II coverage of 99.68% globally and 99.74% in South Asia confirms the potential "
    "for widespread CD4+ helper T-cell activation. Notably, the population-specific coverage for Sri "
    "Lanka was limited by the scarcity of HLA allele frequency data in the IEDB database; the reported "
    "52.39% Class I coverage and the absence of Class II data for Sri Lanka reflect this database "
    "limitation rather than a genuine deficiency in epitope coverage. The broader South Asian coverage "
    "of 99.98%, which geographically encompasses Sri Lanka, supports the expectation that actual Sri "
    "Lankan coverage is substantially higher than the IEDB-computable estimate. Future population-based "
    "HLA typing studies in Sri Lanka, or the incorporation of data from the Allele Frequency Net "
    "Database (AFND), would enable more precise country-level coverage assessment. Collectively, the "
    "population coverage analysis validates the selected epitope ensemble as capable of eliciting both "
    "cellular and humoral immune responses across diverse HLA backgrounds in dengue-endemic populations."
)

doc.add_paragraph("")

# Reference note
p = doc.add_paragraph()
run = p.add_run("Reference to add (if not already in your reference list):")
run.bold = True
run.font.size = Pt(11)
p = doc.add_paragraph()
p.add_run(
    "Bui, H.H., Sidney, J., Dinh, K., Southwood, S., Newman, M.J. and Sette, A., 2006. "
    "Predicting population coverage of T-cell epitope-based diagnostics and vaccines. "
    "BMC Bioinformatics, 7(1), pp.1-5."
).font.size = Pt(10)

doc.save('/workspace/session/Population_Coverage_Results_for_Thesis.docx')
print("DOCX file saved successfully!")

